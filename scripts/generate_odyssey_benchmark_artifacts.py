from __future__ import annotations

import os
from datetime import date, time
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.table import Table, TableStyleInfo


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
DOCX_PATH = OUT_DIR / "OdysseyOne_Requirements_and_User_Stories.docx"
XLSX_PATH = OUT_DIR / "OdysseyOne_Test_Data_and_Mapping.xlsx"

NAVY = "082F49"
BLUE = "2563EB"
TEAL = "0F766E"
PALE_BLUE = "EAF2FF"
PALE_TEAL = "E7F7F3"
PALE_RED = "FDECEC"
LIGHT_GRAY = "F3F6F9"
MID_GRAY = "64748B"
WHITE = "FFFFFF"
BLACK = "0F172A"


LOGIN_STEPS = [
    "Open a fresh browser session.",
    "Navigate to https://qa.linx.odysseylogistics.com/auth/email-classifier?returnUrl=%2Fuser%2Fadministration.",
    "In the Email Address field, enter OdysseyOneAutomationTester1@odysseylogistics.com.",
    "Click the Continue button.",
    "Verify that the Sign in with Microsoft option is displayed.",
    "Click Sign in with Microsoft.",
    "On the Microsoft sign-in page, enter OdysseyOneAutomationTester1@odysseylogistics.com in the email, phone, or Skype field.",
    "Click the Next button.",
    "Enter Behavior-ticket-organize1* in the password field.",
    "Click the Sign in button.",
    "If Microsoft asks whether to stay signed in, choose the option that continues to the application.",
    "Verify that the OdysseyOne Home dashboard is displayed.",
]


ORDER_STEPS = [
    "From the authenticated home dashboard, click the second icon in the left navigation menu, identified as Orders.",
    "Wait for the Orders page to become stable.",
    "Verify that the Orders page is displayed and that the Create Order control is visible and enabled.",
    "Click Create Order.",
    "Wait for the Create New Order page to become stable.",
    "Verify that the Create New Order heading and General Information section are visible.",
    "Fill the Order Number field with 007995145.",
    "Verify that the Order Number field contains exactly 007995145.",
    "Enter SIGROUP, in capital letters, into the Owning Organization field.",
    "Wait until the Owning Organization suggestion list is visible and stable.",
    "Verify that the suggestion list contains these visible options in this order: first *SIGROUP SOURCE SYSTEM 01; second *SIGROUP-EUR SOURCE SYSTEM 01.",
    "Click the second Owning Organization option, *SIGROUP-EUR SOURCE SYSTEM 01.",
    "Verify that the selected Owning Organization is exactly *SIGROUP-EUR SOURCE SYSTEM 01.",
    "Open the Equipment dropdown and wait until its option list is visible and stable.",
    "Verify that the Equipment list contains these visible options in order: RR, LCL, LTL, TL, FCL.",
    "Select the third Equipment option, LTL.",
    "Verify that the Equipment field displays exactly LTL.",
    "Open the Ship Direction dropdown and wait until its options are visible.",
    "Verify that the Ship Direction list contains Outbound and Inbound.",
    "Select Inbound.",
    "Verify that the Ship Direction field displays exactly Inbound.",
    "Wait for any dependent form values to finish updating after the Ship Direction change.",
    "Verify that the Freight Term field has automatically changed from Pre-Paid to exactly COL. If it has not changed to COL, record the expected and observed values as a functional validation failure and continue with the next independent step.",
    "Open the Freight Term dropdown and wait until its options are visible and stable.",
    "Verify that the Freight Term list contains Pre-Paid, Collect, Pre-Paid/Add, Third Party, No Charge, and COL.",
    "Select Collect from the Freight Term dropdown.",
    "Verify that the Freight Term field displays exactly Collect.",
    "Scroll the References section into view.",
    "Fill the Pickup Number field with 7995145776.",
    "Verify that the Pickup Number field contains exactly 7995145776.",
    "Scroll until the Pickup and Delivery section is visible.",
    "Inspect the Pickup and Delivery section. If collapsed, click its header or expand control and wait for it to open. If already expanded, do not click it or collapse it.",
    "Verify that the Pickup and Delivery section is expanded.",
    "Scroll within the expanded Pickup and Delivery section until Planning Date/Time is visible.",
    "Verify that the Planning Date/Time section and the Ship Date & Time option are visible.",
    "Select Ship Date & Time if it is not already selected.",
    "Open the Early Pickup Date calendar and select August 20, 2026.",
    "Verify that Early Pickup Date represents August 20, 2026.",
    "Open the Early Pickup Time dropdown, select 09:00 AM, and verify the selected time.",
    "Open the Early Pickup Time Zone dropdown, select an available option whose visible label contains Central, and verify that the selected label contains Central.",
    "Open the Late Pickup Date calendar and select August 20, 2026.",
    "Verify that Late Pickup Date represents August 20, 2026.",
    "Open the Late Pickup Time dropdown, select 11:00 AM, and verify the selected time.",
    "Open the Late Pickup Time Zone dropdown, select an available option whose visible label contains Central, and verify that the selected label contains Central.",
    "Open the Early Delivery Date calendar and select August 21, 2026.",
    "Verify that Early Delivery Date represents August 21, 2026.",
    "Open the Early Delivery Time dropdown, select 01:00 PM, and verify the selected time.",
    "Open the Early Delivery Time Zone dropdown, select an available option whose visible label contains Central, and verify that the selected label contains Central.",
    "Open the Late Delivery Date calendar and select August 21, 2026.",
    "Verify that Late Delivery Date represents August 21, 2026.",
    "Open the Late Delivery Time dropdown, select 03:00 PM, and verify the selected time.",
    "Open the Late Delivery Time Zone dropdown, select an available option whose visible label contains Central, and verify that the selected label contains Central.",
    "Verify the chronological relationships: Early Pickup Date/Time is before Late Pickup Date/Time; Late Pickup Date/Time is before Early Delivery Date/Time; Early Delivery Date/Time is before Late Delivery Date/Time.",
    "Verify that no required-field validation message is displayed for the completed Planning Date/Time fields.",
    "Leave the completed Create New Order form open without clicking Save, Create Order, or Cancel.",
]


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: str = BLACK, size: float = 9.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], first_width: float = 1.65) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(first_width)
        cells[1].width = Inches(5.7)
        set_cell_text(cells[0], key, bold=True, color=NAVY)
        set_cell_text(cells[1], value)
        shade_cell(cells[0], PALE_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_label(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.color.rgb = RGBColor.from_string(TEAL)
    value = paragraph.add_run(text)
    value.font.name = "Aptos"
    value.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_steps(doc: Document, steps: list[str]) -> None:
    for index, step in enumerate(steps, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(4)
        n = p.add_run(f"{index}. ")
        n.bold = True
        n.font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(step)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(28)
    styles["Title"].font.color.rgb = RGBColor.from_string(NAVY)
    for style_name, size, color in [("Heading 1", 18, NAVY), ("Heading 2", 14, TEAL), ("Heading 3", 11.5, BLUE)]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("OdysseyOne End-to-End Automation Requirements")
    subtitle = doc.add_paragraph()
    subtitle.add_run("QAAI generation benchmark | session continuation | complex control coverage").italic = True
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)

    add_key_value_table(doc, [
        ("Document purpose", "Validate QAAI scenario generation, long-case preservation, data mapping, data binding, and browser-session continuation."),
        ("Expected output", "Exactly 2 scenarios and exactly 2 automation test cases."),
        ("Case relationship", "TC-002 depends on TC-001 and continues in the same authenticated browser session."),
        ("Step preservation", "TC-001 contains 12 authored steps. TC-002 contains 55 authored steps. Do not split, compress, omit, or reorder them."),
        ("Data source", "OdysseyOne_Test_Data_and_Mapping.xlsx; mappings are approved and values must match this document exactly."),
        ("Security", "The workbook contains supplied test credentials and must be handled as sensitive test data."),
    ])

    doc.add_heading("Generation Contract", level=1)
    add_bullets(doc, [
        "Generate exactly two scenarios: Authentication and Create Order.",
        "Generate exactly one test case under each scenario.",
        "Keep the complete 55-step Create Order flow in one test case; it is below the platform maximum of 100 authored steps.",
        "Resolve the second case dependency by the exact first-case name: Login through email classifier and Microsoft sign-in.",
        "TC-001 uses sessionMode=fresh. TC-002 uses sessionMode=continue_from_dependency.",
        "TC-002 must reuse the same browser page, browser context, cookies, and authenticated state produced by TC-001.",
        "Do not navigate directly to the order page or repeat authentication in TC-002.",
        "Persist all authored validations. A non-blocking functional mismatch must not erase later independent checks.",
        "Do not invent values, credentials, substitute dates, labels, options, or assertions.",
    ])

    doc.add_page_break()
    doc.add_heading("US-001 | Email Classifier and Microsoft Sign-In", level=1)
    add_label(doc, "Requirement Title", "OdysseyOne email classifier Microsoft sign-in opens the Home dashboard")
    add_label(doc, "User Story", "As an OdysseyOne internal user, I want to authenticate through the email classifier and Microsoft sign-in so that I can reach the authenticated Home dashboard.")
    add_label(doc, "Target URL", "https://qa.linx.odysseylogistics.com/auth/email-classifier?returnUrl=%2Fuser%2Fadministration")
    add_label(doc, "Scenario", "Verify one continuous authentication flow from the email classifier through Microsoft sign-in to the OdysseyOne Home dashboard.")
    add_label(doc, "Authoring Rule", "Generate exactly one scenario with exactly one test case. Keep the classifier, provider selection, Microsoft credential entry, optional stay-signed-in handling, and final dashboard validation in the same fresh browser session. Do not split this flow.")
    add_label(doc, "Test Case", "Login through email classifier and Microsoft sign-in")

    doc.add_heading("Test Data", level=2)
    add_key_value_table(doc, [
        ("Email Address", "OdysseyOneAutomationTester1@odysseylogistics.com"),
        ("Password", "Behavior-ticket-organize1*"),
        ("Final oracle", 'Home page displayed and text "Welcome OdysseyOne!" visible'),
    ])

    doc.add_heading("Acceptance Criteria", level=2)
    add_bullets(doc, [
        "The email address entered on the classifier page exactly matches the approved AuthProfiles row.",
        "The same email address is reused on the Microsoft sign-in page.",
        "The supplied password is referenced from sensitive test data and is not replaced with a draft value.",
        "An optional stay-signed-in prompt is handled only when present.",
        'The test passes only when the authenticated Home page is displayed and "Welcome OdysseyOne!" is visible.',
    ])

    doc.add_heading("Steps", level=2)
    add_steps(doc, LOGIN_STEPS)

    doc.add_heading("Final Validation", level=2)
    doc.add_paragraph('The preferred final assertion is: verify that the Home page is displayed and the text "Welcome OdysseyOne!" is visible. Supporting dashboard signals may include the OdysseyONE logo, Home, OdysseyOne Automation, Admin, Edit Dashboard View, Order Exceptions, Go to Order, Carriers, What would you like to do?, Go to Create a New Order, Track a Shipment, Management Users, and Invoices.')

    doc.add_heading("Session Policy", level=2)
    add_key_value_table(doc, [
        ("sessionMode", "fresh"),
        ("dependsOnNames", "none"),
        ("failurePolicy", "block_dependents"),
        ("producesState", "authenticated_session; home_dashboard"),
    ])

    doc.add_heading("Data Binding Rule", level=2)
    doc.add_paragraph("Use the approved AuthProfiles row AUTH-001. Preserve the exact email and password. The same email is required in both email-entry controls. Do not invent substitute or placeholder credentials. The workbook duplicates the exact source values so QAAI can prove mapping consistency.")

    doc.add_heading("Expected Scenario/Test Case Shape", level=2)
    add_key_value_table(doc, [
        ("Expected scenarios", "1"),
        ("Expected test cases", "1"),
        ("Expected authored steps", "12"),
        ("Reason", "One coherent authentication behavior with one final business outcome."),
    ])

    doc.add_page_break()
    doc.add_heading("US-002 | Create Order with Complex Controls", level=1)
    add_label(doc, "Requirement Title", "Create an order from the authenticated dashboard and validate complex form controls")
    add_label(doc, "User Story", "As an authenticated OdysseyOne operations user, I want to populate and verify a Create New Order form so that complex dropdown, autocomplete, expandable-section, calendar, time, time-zone, and dependent-field behavior is validated without submitting the order.")
    add_label(doc, "Scenario", "Continue from the authenticated Home dashboard, open Create New Order, populate General Information, References, and Planning Date/Time, verify every value and option contract, and leave the form open.")
    add_label(doc, "Authoring Rule", "Generate exactly one scenario with exactly one 55-step test case. Preserve every numbered step, value, assertion, conditional instruction, and continuation rule. Do not split or compress this case. Inline values must remain explicit in the authored steps and must match the approved workbook row.")
    add_label(doc, "Test Case", "Create an order from the authenticated dashboard and validate complex form controls")

    doc.add_heading("Initial State and Dependency", level=2)
    add_bullets(doc, [
        "The previous test case, Login through email classifier and Microsoft sign-in, completed successfully.",
        "The current browser session is authenticated and remains on the OdysseyOne Home dashboard.",
        "Reuse the same browser page, context, cookies, and authenticated session.",
        "Do not launch a new browser, create a new context, navigate directly by URL, or perform login again.",
    ])

    doc.add_heading("Inline Test Data", level=2)
    add_key_value_table(doc, [
        ("Order Number", "007995145"),
        ("Owning Org search", "SIGROUP"),
        ("Owning Org selection", "*SIGROUP-EUR SOURCE SYSTEM 01"),
        ("Equipment", "LTL"),
        ("Ship Direction", "Inbound"),
        ("Automatic Freight Term", "COL"),
        ("Final Freight Term", "Collect"),
        ("Pickup Number", "7995145776"),
        ("Early Pickup", "08/20/2026 09:00 AM"),
        ("Late Pickup", "08/20/2026 11:00 AM"),
        ("Early Delivery", "08/21/2026 01:00 PM"),
        ("Late Delivery", "08/21/2026 03:00 PM"),
        ("Time Zone rule", 'Select an available option whose visible label contains "Central"'),
    ])

    doc.add_heading("Steps and Validations", level=2)
    add_steps(doc, ORDER_STEPS)

    doc.add_heading("Failure and Continuation Behavior", level=2)
    add_bullets(doc, [
        "Wait for every dropdown, autocomplete list, or calendar surface before selecting.",
        "Verify every selected value from the live control after selection.",
        "When expected list order or displayed value is wrong, record the exact expected and observed values as a functional validation failure and continue with independent checks.",
        "When a required control cannot be uniquely resolved or an action cannot be confirmed, perform one fresh-page evidence retry, record a QAAI execution uncertainty if still unresolved, and stop only dependent steps.",
        "Do not report a website failure when browser evidence is inconclusive.",
        "The COL auto-update check is non-blocking for the later explicit Collect selection and remaining independent validations.",
    ])

    doc.add_heading("Expected Final State", level=2)
    doc.add_paragraph("The Create New Order form remains open with the specified General Information, References, and Planning Date/Time values populated and verified. Do not click Save, Create Order, or Cancel.")

    doc.add_heading("Session Policy", level=2)
    add_key_value_table(doc, [
        ("sessionMode", "continue_from_dependency"),
        ("dependsOnNames", "Login through email classifier and Microsoft sign-in"),
        ("requiresState", "authenticated_session; home_dashboard"),
        ("failurePolicy", "block_dependents when the prerequisite fails; continue independent in-case checks after non-blocking validation mismatch"),
    ])

    doc.add_heading("Data Binding Rule", level=2)
    doc.add_paragraph("Use approved row ORDER-001 from OrderCreation and the linked OptionExpectations rows. Keep every inline value visible in the generated test steps and verify it against the mapped workbook value. Do not invent or substitute values. Date cells are typed Excel dates, and time cells are typed Excel times.")

    doc.add_heading("Expected Scenario/Test Case Shape", level=2)
    add_key_value_table(doc, [
        ("Expected scenarios", "1"),
        ("Expected test cases", "1"),
        ("Expected authored steps", "55"),
        ("Reason", "One stateful form-completion behavior that must preserve the authenticated browser session and all dependent control state."),
    ])

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Cross-Artifact Validation Checklist", level=1)
    add_bullets(doc, [
        "Scenario count equals 2 and test-case count equals 2.",
        "TC-001 has exactly 12 authored steps and sessionMode=fresh.",
        "TC-002 has exactly 55 authored steps and sessionMode=continue_from_dependency.",
        "TC-002 depends on TC-001 by exact case name and does not repeat login.",
        "Every supplied value is present in the generated case and matches the approved workbook mapping.",
        "Dropdown option order assertions are retained.",
        "Conditional expand/no-collapse behavior is retained.",
        "The automatic COL mismatch is recorded without preventing later independent steps.",
        "The form remains open and no final submit/cancel action is generated.",
    ])

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


THIN = Side(style="thin", color="D9E2EC")


def style_sheet(ws, widths: dict[str, float] | None = None, freeze: str = "A2") -> None:
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = freeze
    ws.auto_filter.ref = ws.dimensions
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(name="Aptos", size=10, bold=True, color=WHITE)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=THIN)
    ws.row_dimensions[1].height = 30
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = Font(name="Aptos", size=9.5, color=BLACK)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(bottom=THIN)
    if widths:
        for col, width in widths.items():
            ws.column_dimensions[col].width = width


def add_table(ws, name: str) -> None:
    if ws.max_row < 2 or ws.max_column < 1:
        return
    table = Table(displayName=name, ref=ws.dimensions)
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    ws.add_table(table)


def append_rows(ws, headers: list[str], rows: list[list[object]]) -> None:
    ws.append(headers)
    for row in rows:
        ws.append(row)


def build_xlsx() -> None:
    wb = Workbook()
    readme = wb.active
    readme.title = "README"
    append_rows(readme, ["item", "value"], [
        ["Purpose", "Approved test data, semantic mappings, option expectations, and session contracts for the two-story QAAI benchmark."],
        ["Source authority", "Uploaded requirement document > approved workbook mapping > browser evidence > inference."],
        ["Expected generation", "2 scenarios; 2 test cases; 12 login steps; 55 order steps."],
        ["Sensitive data", "AuthProfiles contains the user-supplied test password in plaintext. Restrict access and do not commit this workbook."],
        ["Binding mode", "Mappings are approved. Exact values must be preserved; no first-row fallback and no invented substitutes."],
        ["Session rule", "TC-002 continues from TC-001 in the same page/context/cookies/authenticated session."],
        ["Dates and times", "OrderCreation date cells are typed dates and time cells are typed times, with display formats applied."],
    ])
    style_sheet(readme, {"A": 24, "B": 105})
    add_table(readme, "ReadmeTable")

    auth = wb.create_sheet("AuthProfiles")
    append_rows(auth, [
        "auth_profile_id", "role", "email_address", "password", "provider_type", "target_url",
        "post_login_oracle", "session_mode", "mapping_status", "sensitive",
    ], [[
        "AUTH-001", "internal_user", "OdysseyOneAutomationTester1@odysseylogistics.com",
        "Behavior-ticket-organize1*", "sso", "https://qa.linx.odysseylogistics.com/auth/email-classifier?returnUrl=%2Fuser%2Fadministration",
        "Welcome OdysseyOne!", "fresh", "approved", True,
    ]])
    style_sheet(auth, {"A": 18, "B": 18, "C": 44, "D": 34, "E": 18, "F": 88, "G": 28, "H": 18, "I": 18, "J": 14})
    add_table(auth, "AuthProfilesTable")
    auth["D2"].fill = PatternFill("solid", fgColor=PALE_RED)

    order = wb.create_sheet("OrderCreation")
    headers = [
        "case_data_id", "order_number", "owning_org_search", "owning_org_selection", "equipment",
        "ship_direction", "expected_auto_freight_term", "final_freight_term", "pickup_number",
        "early_pickup_date", "early_pickup_time", "early_pickup_timezone_contains",
        "late_pickup_date", "late_pickup_time", "late_pickup_timezone_contains",
        "early_delivery_date", "early_delivery_time", "early_delivery_timezone_contains",
        "late_delivery_date", "late_delivery_time", "late_delivery_timezone_contains",
        "submit_order", "mapping_status",
    ]
    append_rows(order, headers, [[
        "ORDER-001", "007995145", "SIGROUP", "*SIGROUP-EUR SOURCE SYSTEM 01", "LTL",
        "Inbound", "COL", "Collect", "7995145776",
        date(2026, 8, 20), time(9, 0), "Central",
        date(2026, 8, 20), time(11, 0), "Central",
        date(2026, 8, 21), time(13, 0), "Central",
        date(2026, 8, 21), time(15, 0), "Central",
        False, "approved",
    ]])
    style_sheet(order, {chr(64 + i): 22 for i in range(1, 24)})
    order.column_dimensions["D"].width = 38
    add_table(order, "OrderCreationTable")
    for col in ["J", "M", "P", "S"]:
        order[f"{col}2"].number_format = "mm/dd/yyyy"
    for col in ["K", "N", "Q", "T"]:
        order[f"{col}2"].number_format = "hh:mm AM/PM"

    options = wb.create_sheet("OptionExpectations")
    append_rows(options, ["case_data_id", "control", "sequence", "visible_label", "required", "mapping_status"], [
        ["ORDER-001", "owning_organization", 1, "*SIGROUP SOURCE SYSTEM 01", True, "approved"],
        ["ORDER-001", "owning_organization", 2, "*SIGROUP-EUR SOURCE SYSTEM 01", True, "approved"],
        ["ORDER-001", "equipment", 1, "RR", True, "approved"],
        ["ORDER-001", "equipment", 2, "LCL", True, "approved"],
        ["ORDER-001", "equipment", 3, "LTL", True, "approved"],
        ["ORDER-001", "equipment", 4, "TL", True, "approved"],
        ["ORDER-001", "equipment", 5, "FCL", True, "approved"],
        ["ORDER-001", "ship_direction", 1, "Outbound", True, "approved"],
        ["ORDER-001", "ship_direction", 2, "Inbound", True, "approved"],
        ["ORDER-001", "freight_term", 1, "Pre-Paid", True, "approved"],
        ["ORDER-001", "freight_term", 2, "Collect", True, "approved"],
        ["ORDER-001", "freight_term", 3, "Pre-Paid/Add", True, "approved"],
        ["ORDER-001", "freight_term", 4, "Third Party", True, "approved"],
        ["ORDER-001", "freight_term", 5, "No Charge", True, "approved"],
        ["ORDER-001", "freight_term", 6, "COL", True, "approved"],
    ])
    style_sheet(options, {"A": 18, "B": 28, "C": 12, "D": 42, "E": 12, "F": 18})
    add_table(options, "OptionExpectationsTable")

    mappings = wb.create_sheet("FieldMapping")
    mapping_rows = [
        ["AuthProfiles", "email_address", "auth.email", "TC-001", "fill_classifier_and_provider_email", True, True, "approved"],
        ["AuthProfiles", "password", "auth.password", "TC-001", "fill_provider_password", True, True, "approved"],
        ["AuthProfiles", "target_url", "navigation.start_url", "TC-001", "initial_navigation", True, False, "approved"],
        ["AuthProfiles", "post_login_oracle", "oracle.dashboard_text", "TC-001", "final_assertion", True, False, "approved"],
    ]
    for column in headers[1:21]:
        mapping_rows.append(["OrderCreation", column, f"order.{column}", "TC-002", "fill_or_verify_exact_value", True, False, "approved"])
    mapping_rows.extend([
        ["OptionExpectations", "visible_label", "control.option_label", "TC-002", "ordered_list_assertion", True, False, "approved"],
        ["OptionExpectations", "sequence", "control.option_sequence", "TC-002", "ordered_list_assertion", True, False, "approved"],
    ])
    append_rows(mappings, ["source_sheet", "source_column", "semantic_role", "test_case_id", "usage", "required", "sensitive", "mapping_status"], mapping_rows)
    style_sheet(mappings, {"A": 22, "B": 34, "C": 38, "D": 16, "E": 36, "F": 12, "G": 12, "H": 18})
    add_table(mappings, "FieldMappingTable")

    sessions = wb.create_sheet("SessionContracts")
    append_rows(sessions, [
        "test_case_id", "test_case_name", "session_mode", "depends_on_test_case_id", "depends_on_test_case_name",
        "requires_state", "produces_state", "failure_policy", "initial_state", "expected_final_state",
    ], [
        ["TC-001", "Login through email classifier and Microsoft sign-in", "fresh", "", "", "", "authenticated_session;home_dashboard", "block_dependents", "fresh browser", "authenticated Home dashboard"],
        ["TC-002", "Create an order from the authenticated dashboard and validate complex form controls", "continue_from_dependency", "TC-001", "Login through email classifier and Microsoft sign-in", "authenticated_session;home_dashboard", "populated_unsaved_order_form", "block_dependents", "same authenticated page/context/cookies", "Create New Order form populated and left open"],
    ])
    style_sheet(sessions, {"A": 14, "B": 58, "C": 28, "D": 24, "E": 58, "F": 34, "G": 34, "H": 22, "I": 42, "J": 52})
    add_table(sessions, "SessionContractsTable")

    expected = wb.create_sheet("ExpectedGeneration")
    append_rows(expected, [
        "scenario_id", "scenario_title", "expected_case_count", "test_case_id", "test_case_name",
        "expected_step_count", "session_mode", "depends_on", "split_allowed", "notes",
    ], [
        ["TS-001", "Email classifier Microsoft sign-in to Home dashboard", 1, "TC-001", "Login through email classifier and Microsoft sign-in", 12, "fresh", "", False, "One coherent authentication flow"],
        ["TS-002", "Create order and validate complex form controls", 1, "TC-002", "Create an order from the authenticated dashboard and validate complex form controls", 55, "continue_from_dependency", "TC-001", False, "Preserve all 55 steps in one case; do not submit the order"],
    ])
    style_sheet(expected, {"A": 16, "B": 52, "C": 20, "D": 16, "E": 62, "F": 22, "G": 28, "H": 18, "I": 18, "J": 56})
    add_table(expected, "ExpectedGenerationTable")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    wb.save(XLSX_PATH)


def validate_artifacts() -> None:
    doc = Document(DOCX_PATH)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Login through email classifier and Microsoft sign-in" in text
    assert "Create an order from the authenticated dashboard and validate complex form controls" in text
    assert "55 authored steps" in text
    for expected in ["007995145", "7995145776", "Welcome OdysseyOne!", "continue_from_dependency"]:
        assert expected in text

    wb = load_workbook(XLSX_PATH, data_only=False)
    assert set(["README", "AuthProfiles", "OrderCreation", "OptionExpectations", "FieldMapping", "SessionContracts", "ExpectedGeneration"]).issubset(wb.sheetnames)
    assert wb["ExpectedGeneration"]["F3"].value == 55
    assert wb["SessionContracts"]["C3"].value == "continue_from_dependency"
    assert wb["AuthProfiles"]["C2"].value == "OdysseyOneAutomationTester1@odysseylogistics.com"
    assert wb["OrderCreation"]["B2"].value == "007995145"
    assert len(ORDER_STEPS) == 55
    assert len(LOGIN_STEPS) == 12


if __name__ == "__main__":
    build_docx()
    build_xlsx()
    validate_artifacts()
    print(DOCX_PATH)
    print(XLSX_PATH)
    print(f"login_steps={len(LOGIN_STEPS)} order_steps={len(ORDER_STEPS)}")
